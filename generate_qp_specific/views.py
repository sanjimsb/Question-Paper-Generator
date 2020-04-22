# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from django.views.generic import TemplateView
from django.shortcuts import render, redirect
from . forms import g_qp_specific
from add_question_paper.models import question_details
from django.core import serializers
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint
import random
# Create your views here.

class choose_question(TemplateView):

	template_name = 'generate_qp_specific/choose_question.html'

	def get_value(self):
		return self.request.GET.get('sub')

	def get(self,request,*args,**kwargs):
		return render(request,self.template_name)

	def post(self,request,*args,**kwargs):
		if request.method == 'POST':
			if 'csb' in request.POST:
				subject = request.POST['subject']
				clickable_option_a = request.POST['g_q_a']
				clickable_option_b = request.POST['g_q_b']
				clickable_option_c = request.POST['g_q_c']
				choosable_questions_a = request.POST['g_a_cq']
				choosable_questions_b = request.POST['g_b_cq']
				choosable_questions_c = request.POST['g_c_cq']
				institute_name = request.POST['iname']
				term = request.POST['term']
				level = request.POST['class']
				time =  request.POST['time']


				request.session['lvl'] = level

				request.session['tm'] = time

				request.session['term_opt'] = term

				request.session['in'] = institute_name

				request.session['gqa'] = clickable_option_a
				request.session['gqb'] = clickable_option_b
				request.session['gqc'] = clickable_option_c

				request.session['gacq'] = choosable_questions_a
				request.session['gbcq'] = choosable_questions_b
				request.session['gccq'] = choosable_questions_c


				request.session['sub'] = subject
				
		return redirect('http://127.0.0.1:8000/generate_qp_specific/')
		



class qp_choice_generate(TemplateView):
	
	template_name = 'generate_qp_specific/index.html'

	def get(self,request,*args,**kwargs):
		form = g_qp_specific()
		question_lst = self.question_option()
		group_a_ques = question_lst[0]
		group_b_ques = question_lst[1]
		group_c_ques = question_lst[2]

		num_question = self.get_subject()
		groupa_selection_ques =  num_question[1]
		groupb_selection_ques =  num_question[2]
		groupc_selection_ques =  num_question[3]

		limit_selection = {
				'groupa_selection_ques':groupa_selection_ques,
				'groupb_selection_ques':groupb_selection_ques,
				'groupc_selection_ques':groupc_selection_ques
		}

		return render(request,self.template_name,{'form':form,'group_a_ques':group_a_ques, 'group_b_ques':group_b_ques, 'group_c_ques':group_c_ques,'limit_selection':limit_selection})
		

	def post(self,request,*args,**kwargs):
		form = g_qp_specific(request.POST)
		if request.method == 'POST':		
			questions_to_ceate_doc = self.getquestions()
			group_a_ques = questions_to_ceate_doc[0]
			group_b_ques = questions_to_ceate_doc[1]
			group_c_ques = questions_to_ceate_doc[2]

			choosable_questions = self.choosable_questions()

			lst_que = {
					'group_a_ques':group_a_ques,				
					'group_b_ques':group_b_ques,				
					'group_c_ques':group_c_ques,
					'choosable_questions':choosable_questions
			}
			self.generate_word(**lst_que)

			sub = self.get_subject()
			chose_sub = sub[0]
		
		return redirect('http://127.0.0.1:8000/media/%s_question.docx'%chose_sub)
	

	def get_subject(self):
		return self.request.session['sub'], self.request.session['gqa'], self.request.session['gqb'], self.request.session['gqc']


	def choosable_questions(self):
		return self.request.session['gacq'], self.request.session['gbcq'], self.request.session['gccq'], self.request.session['in']


	def fetch_question_group(self):
		subject = self.get_subject()
		sub = subject[0]
		group_a_questions = question_details.objects.filter(subject = sub, question_weight = 2 )
		group_b_questions = question_details.objects.filter(subject = sub, question_weight = 5 )
		group_c_questions = question_details.objects.filter(subject = sub, question_weight = 10 )

		return group_a_questions, group_b_questions, group_c_questions


	def question_option(self):
		get_subject = self.get_subject()
		sub = get_subject[0]

		questions_number = self.fetch_question_group()
		group_a_questions = questions_number[0]
		group_b_questions = questions_number[1]
		group_c_questions = questions_number[2]

		group_a_total_questions = len(group_a_questions)
		group_a_questions_list = []

		group_b_total_questions = len(group_b_questions)
		group_b_questions_list = []

		group_c_total_questions = len(group_c_questions)
		group_c_questions_list = []

		for i in range(0 ,group_a_total_questions):
			group_a_questions_list.append(group_a_questions[i].question)

		for i in range(0 ,group_b_total_questions):
			group_b_questions_list.append(group_b_questions[i].question)

		for i in range(0 ,group_c_total_questions):
			group_c_questions_list.append(group_c_questions[i].question)

		return group_a_questions_list, group_b_questions_list, group_c_questions_list


	def getquestions(self):
		get_subject = self.fetch_question_group()


		groupa_length = int(len(get_subject[0]))
		selected_question_values_a = []
		try:
			for i in range(1,groupa_length + 1):
				ga_questions = self.request.POST.get('group_a%s'%i)
				selected_question_values_a.append(ga_questions)
		except Exception as e:
			print('error')


		groupb_length = int(len(get_subject[1]))
		selected_question_values_b = []
		try:
			for i in range(1,groupb_length + 1):
				gb_questions = self.request.POST.get('group_b%s'%i)
				selected_question_values_b.append(gb_questions)
		except Exception as e:
			print('error')


		groupc_length = int(len(get_subject[2]))
		selected_question_values_c = []
		try:
			for i in range(1,groupc_length + 1):
				gc_questions = self.request.POST.get('group_c%s'%i)
				selected_question_values_c.append(gc_questions)
		except Exception as e:
			print('error')

		selected_question_values_a	=	filter(None,selected_question_values_a)
		selected_question_values_b	=	filter(None,selected_question_values_b)
		selected_question_values_c	=	filter(None,selected_question_values_c)


		return selected_question_values_a, selected_question_values_b, selected_question_values_c  


	def generate_word(self,*args,**kwargs):
		
		questions = kwargs
		group_a_q = questions["group_a_ques"]
		group_b_q = questions["group_b_ques"]
		group_c_q = questions["group_c_ques"]
		choosable_option = questions["choosable_questions"]

		ga = int(choosable_option[0])
		gb = int(choosable_option[1])
		gc = int(choosable_option[2])

		sub = self.get_subject()
		choose_term = self.request.session['term_opt']
		time = self.request.session['tm']
		level = self.request.session['lvl']

		total_marks = ((ga * 2) + (gb * 5) + (gc * 10))
		pass_marks = int((total_marks * 32) / 100)

		document = Document()

		heading = document.add_heading('%s'%choosable_option[3], 2)
		heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

		term = document.add_paragraph('%s'%(choose_term))
		term.alignment = WD_ALIGN_PARAGRAPH.CENTER


		fullmarks = document.add_paragraph('Subject : %s                                                                                                                   Full marks :%d'%(sub[0] ,total_marks))
		passmarks = document.add_paragraph('Class :%s                                                                                                                                   Pass marks : %d'%(level, pass_marks))
		add_time = document.add_paragraph('Time :%s'%time)
		add_time.alignment = WD_ALIGN_PARAGRAPH.LEFT
		fullmarks.alignment = WD_ALIGN_PARAGRAPH.LEFT
		passmarks.alignment = WD_ALIGN_PARAGRAPH.LEFT

		# prior_paragraph = fullmarks.insert_paragraph_before('Subject:')
		# prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

		space = document.add_paragraph()
		group_fst = document.add_paragraph()
		group_fst.add_run('Group A').bold = True
		group_fst.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_fst_heading = document.add_paragraph('Answer the following question(choose any %s)'%(choosable_option[0]))
		group_fst_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		for count, i in enumerate(group_a_q, start = 1):
			question_ga = document.add_paragraph('%d. %s'%(count, i))


		space = document.add_paragraph()	
		group_second = document.add_paragraph()
		group_second.add_run('Group B').bold = True
		group_second.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_second_heading = document.add_paragraph('Answer the following question(choose any %s)'%(choosable_option[1]))
		group_second_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		for count, i in enumerate(group_b_q, start = 1):
			question_gb = document.add_paragraph('%d. %s'%(count, i))


		space = document.add_paragraph()
		group_third = document.add_paragraph()
		group_third.add_run('Group C').bold = True
		group_third.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_third_heading = document.add_paragraph('Answer the following question(choose any %s)'%(choosable_option[2]))
		group_third_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		for count, i in enumerate(group_c_q,start = 1):
			question_gc = document.add_paragraph('%d. %s'%(count, i))

		space = document.add_paragraph()
		thanks = document.add_paragraph()
		thanks.add_run('BEST OF LUCK').bold=True
		thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER


		document.save('media/Account_question.docx')
		return None



















# class qp_choice_generate(TemplateView):
	
# 	template_name = 'generate_qp_specific/index.html'

# 	def get(self,request,*args,**kwargs):
# 		form = g_qp_specific()
# 		if self.request.is_ajax():
# 			return self.ajax(request)
# 		return None
		
		

# 	def post(self,request,*args,**kwargs):
# 		form = g_qp_specific(request.POST)
# 		sub = self.get_subject()
# 		print sub
# 		return render(request,self.template_name,{'form':form})
	
# 	def get_subject(self):
# 		return self.request.GET.get('sub')

# 	def ajax(self,request):
# 		sub = self.get_subject()
# 		a = question_details.objects.filter(subject=sub)
# 		b = len(a)
# 		c = []
# 		print a 
# 		print sub
# 		for i in range(0 ,b - 1):
# 			c.append(a[i].question)

# 		question_lst = c
# 		print question_lst
# 		# non_ser = question_lst[0]
# 		# json_serializer = serializers.get_serializer("json")()
# 		# serialized_info = json_serializer.serialize(question_lst[1])
# 		# aaa = len(serialized_info)
# 		# print aaa
# 		return render(request,self.template_name,{'question_lst':question_lst})



