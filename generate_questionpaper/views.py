# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from django.views.generic import TemplateView
from django.shortcuts import render, redirect
from django.http import HttpResponseRedirect
from . forms import generate_q
from . models import generate_options
from add_question_paper.models import question_details
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from random import randint
import random
from django.core.mail import send_mail
# Create your views here.


class generate_question_paper(TemplateView):
	template_name = 'generate_questionpaper/index.html'


	def generate_word(self,*args,**kwargs):
		data = kwargs
		document = Document()

		ga = int(data["Group_A_required_questions"])
		gb = int(data["group_b_required_questions"])
		gc = int(data["group_c_required_questions"])
		total_marks = ((ga * 2) + (gb * 5) + (gc * 10))
		pass_marks = (total_marks * 32) / 100

		heading = document.add_heading(data["institute_name"], 2)
		heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

		fullmarks = document.add_paragraph('Subject : %s                                                                                                                     Full marks :%s'%(data["Subject"], total_marks))
		passmarks = document.add_paragraph('Class : %s                                                                                                                                   Pass marks :%s '%(data["level"], pass_marks))
		fullmarks.alignment = WD_ALIGN_PARAGRAPH.LEFT
		passmarks.alignment = WD_ALIGN_PARAGRAPH.LEFT

		# prior_paragraph = fullmarks.insert_paragraph_before('Subject:')
		# prior_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
		space = document.add_paragraph()
		group_fst = document.add_paragraph('Group A')
		group_fst.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_fst_heading = document.add_paragraph('Answer the following question(choose any %d)' %int(data["Group_A_required_questions"]))
		group_fst_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		group_a_questions = question_details.objects.filter(subject = data["Subject"], question_weight = 2)
		subject_tally_a = group_a_questions.count()

		group_a_all_q = data["Group_A_questions"]
		group_a_q_len = int(group_a_all_q)

		length_a = len(group_a_questions)
		new_list_a = []
		for i in range(0,length_a):
			new_list_a.append(group_a_questions[i].question)

		for i in range(1,group_a_q_len + 1):
			random.shuffle(new_list_a)
			question_ga = document.add_paragraph('%d. %s'%(i, new_list_a.pop()))



		space = document.add_paragraph()
		group_second = document.add_paragraph('Group B')
		group_second.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_second_heading = document.add_paragraph('Answer the following question(choose any %d)' %int(data["group_b_required_questions"]))
		group_second_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		group_b_questions = question_details.objects.filter(subject = data["Subject"], question_weight = 5)
		subject_tally_b = group_b_questions.count()
		
		group_b_all_q = data["group_b_questions"]
		group_b_q_len = int(group_b_all_q)

		length_b = len(group_b_questions)
		new_list_b = []
		
		for i in range(0,length_b):
			new_list_b.append(group_b_questions[i].question)


		for i in range(1, group_b_q_len + 1):
			random.shuffle(new_list_b)
			question_gb = document.add_paragraph('%d. %s'%(i, new_list_b.pop()))



		space = document.add_paragraph()
		group_third = document.add_paragraph('Group C')
		group_third.alignment = WD_ALIGN_PARAGRAPH.CENTER

		group_third_heading = document.add_paragraph('Answer the following question(choose any %d)' %int(data["group_c_required_questions"]))
		group_third_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

		group_c_questions = question_details.objects.filter(subject = data["Subject"], question_weight = 10)
		subject_tally_c = group_c_questions.count()
		
		group_c_all_q = data["group_c_questions"]
		group_c_q_len = int(group_c_all_q)

		length_c = len(group_c_questions)
		new_list_c = []
		
		for i in range(0,length_c):
			new_list_c.append(group_c_questions[i].question)


		for i in range(1, group_c_q_len + 1):
			random.shuffle(new_list_c)
			question_gc = document.add_paragraph('%d. %s'%(i, new_list_c.pop()))

		document.save('media/%s_question.docx' %data["Subject"])
		return None


	def get(self,request,*args,**kwargs):
		form = generate_q()
		return render(request,self.template_name,{'form':form})


	def post(self,request,*args,**kwargs):
		form = generate_q(request.POST)	
		if form.is_valid():
			if 'generate' in request.POST:
				data = {
					'institute_name'				: form.cleaned_data['institute_name'],	
					'Subject'						: form.cleaned_data['Subject'],
					'level'							: form.cleaned_data['Level'],
					'Group_A_questions'				: form.cleaned_data['Group_A_questions'],			
					'Group_A_required_questions'	: form.cleaned_data['Group_A_required_questions'],			
					'group_b_questions'				: form.cleaned_data['group_b_questions'],			
					'group_b_required_questions'	: form.cleaned_data['group_b_required_questions'],			
					'group_c_questions'				: form.cleaned_data['group_c_questions'],	
					'group_c_required_questions'	: form.cleaned_data['group_c_required_questions']			
				}
				self.generate_word(**data)
				return redirect("http://127.0.0.1:8000/media/%s_question.docx" %(data["Subject"]))

					
		return render(request,self.template_name,{'form':form})

